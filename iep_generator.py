
import streamlit as st
from docx import Document
from io import BytesIO

st.set_page_config(page_title="IEP 智慧填寫小幫手")

st.title("📄 IEP 個別化教育計畫產出工具")

st.header("基本資料")
student_name = st.text_input("學生姓名")
gender = st.radio("性別", ["男", "女"])
grade = st.text_input("年級")
class_name = st.text_input("班級")
birth_date = st.date_input("出生日期")
meeting_date = st.date_input("開會日期")

st.header("🧠 能力與需求")
attention_issues = st.checkbox("注意力短暫")
memory_issues = st.checkbox("記憶困難")
expression_issues = st.checkbox("表達能力弱")
behavior_note = st.text_area("補充說明", height=100)

st.header("📚 教學策略")
multisensory = st.checkbox("多感官教學")
direct_instruction = st.checkbox("直接教學")
memory_strategy = st.checkbox("記憶策略")

st.header("🎯 學期學習目標")
term_goal = st.text_area("請輸入學期目標", height=100)
evaluation_method = st.multiselect("評量方式", ["書面（紙筆）", "口語問答", "觀察", "操作（實作）"])

if st.button("📝 產出 IEP 草稿"):
    doc = Document()
    doc.add_heading("IEP 個別化教育計畫草稿", level=1)
    doc.add_paragraph(f"學生姓名：{student_name}　性別：{gender}　年級：{grade}　班級：{class_name}")
    doc.add_paragraph(f"出生日期：{birth_date}　開會日期：{meeting_date}")
    doc.add_heading("學生能力與需求", level=2)
    if attention_issues:
        doc.add_paragraph("學生注意力短暫，需給予額外提示與引導。")
    if memory_issues:
        doc.add_paragraph("學生記憶力需加強，建議使用記憶策略。")
    if expression_issues:
        doc.add_paragraph("學生表達能力弱，建議多提供口語練習機會。")
    if behavior_note:
        doc.add_paragraph("補充說明：" + behavior_note)
    doc.add_heading("教學策略", level=2)
    strategies = []
    if multisensory:
        strategies.append("採用多感官教學（如視覺、聽覺、操作）")
    if direct_instruction:
        strategies.append("使用直接教學法")
    if memory_strategy:
        strategies.append("加入記憶策略協助學習")
    if strategies:
        doc.add_paragraph("教學建議：" + "，".join(strategies) + "。")
    doc.add_heading("學習目標與評量", level=2)
    doc.add_paragraph("學期目標：" + term_goal)
    if evaluation_method:
        doc.add_paragraph("評量方式：" + "、".join(evaluation_method))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📥 下載 Word 文件",
        data=buffer,
        file_name=f"{student_name}_IEP.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
